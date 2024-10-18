import fitz  # PyMuPDF for PDF extraction
import docx  # Python-docx for DOCX extraction
import PyPDF2  # For PDF link extraction using text extraction
import re  # For regular expression operations

class LinksExtractor:
    def __init__(self, file_path):
        self.file_path = file_path

    def extract_links_from_pdf_with_fitz(self):
        """
        Extract all the links (URIs) from a PDF file using PyMuPDF.
        """
        links = []
        try:
            # Open the PDF document
            doc = fitz.open(self.file_path)
            
            # Loop through each page of the PDF
            for i in range(doc.page_count):
                page = doc.load_page(i)  # Load each page
                page_links = page.get_links()  # Get all the links on the page

                # Extract and append the URIs to the list
                for link in page_links:
                    if 'uri' in link:
                        links.append(link['uri'])  # Extract the URL if it's a link
        except Exception as e:
            print(f"Error extracting links from PDF using PyMuPDF: {e}")
        return links

    def extract_links_from_pdf_with_pypdf2(self):
        """
        Extract all the links (URIs) from a PDF file using PyPDF2.
        """
        links = []
        try:
            with open(self.file_path, 'rb') as file:
                reader = PyPDF2.PdfReader(file)
                for page_num in range(len(reader.pages)):
                    page = reader.pages[page_num]
                    text = page.extract_text()
                    links += re.findall(r'(https?://\S+)', text)  # Use regex to find URLs
        except Exception as e:
            print(f"Error extracting links from PDF using PyPDF2: {e}")
        return links

    def extract_links_from_docx_with_hyperlink(self):
        """
        Extract all the links (URIs) from a DOCX file using hyperlinks.
        """
        links = []
        try:
            # Open the DOCX document
            doc = docx.Document(self.file_path)
            
            # Loop through each paragraph and extract hyperlinks
            for paragraph in doc.paragraphs:
                for run in paragraph.runs:
                    if run.hyperlink:  # Check if the run has a hyperlink
                        links.append(run.hyperlink.target)  # Extract the URL
        except Exception as e:
            print(f"Error extracting links from DOCX using hyperlinks: {e}")
        return links

    def extract_links_from_docx_with_text(self):
        """
        Extract all the links (URIs) from a DOCX file using text extraction.
        """
        links = []
        try:
            doc = docx.Document(self.file_path)
            for para in doc.paragraphs:
                text = para.text
                links += re.findall(r'(https?://\S+)', text)  # Use regex to find URLs
        except Exception as e:
            print(f"Error extracting links from DOCX using text extraction: {e}")
        return links

    def extract_links(self):
        """
        Extract links from supported file types (PDF and DOCX).
        This method is designed to allow the user to choose which extraction method to use.
        """
        if self.file_path.endswith('.pdf'):
            # You can choose which extraction method to use here:
            return (self.extract_links_from_pdf_with_fitz() + 
                    self.extract_links_from_pdf_with_pypdf2())
        elif self.file_path.endswith('.docx'):
            # You can choose which extraction method to use here:
            return (self.extract_links_from_docx_with_hyperlink() + 
                    self.extract_links_from_docx_with_text())
        else:
            print("Unsupported file format. Please provide a .pdf or .docx file.")
            return []
